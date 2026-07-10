# -*- coding: utf-8 -*-
"""Convierte manuscrito_estudios_economia.md en el .docx de envío para
Estudios de Economía (12 pt, doble espacio, márgenes amplios, tablas
editables con fuente).

Requiere: python-docx.  Uso: python generar_docx_epp.py
"""
import os
import re
import sys

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
import latex2mathml.converter
from lxml import etree

_XSL_OMML = etree.XSLT(etree.parse(
    r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL'))

def latex_a_omml(latex):
    """Convierte LaTeX a un elemento OMML (ecuacion nativa de Word)."""
    mathml = latex2mathml.converter.convert(latex)
    omml = _XSL_OMML(etree.fromstring(mathml))
    return omml.getroot()

AQUI = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(AQUI, 'manuscrito_estudios_economia.md')
OUT = os.path.join(AQUI, 'manuscrito_estudios_economia.docx')
DIR_PAPER = os.path.dirname(AQUI)  # articulo/

doc = Document()

# ---------------- Estilos base ----------------
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

for sec in doc.sections:
    sec.top_margin = Cm(3)
    sec.bottom_margin = Cm(3)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(3)

# ---------------- Utilidades ----------------
TOKEN = re.compile(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+?`)')

def limpiar(s):
    s = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', s)
    s = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', s)
    return s

def agregar_runs(p, texto, tamano=None):
    texto = limpiar(texto)
    for parte in TOKEN.split(texto):
        if not parte:
            continue
        run = None
        if parte.startswith('***') and parte.endswith('***'):
            run = p.add_run(parte[3:-3]); run.bold = True; run.italic = True
        elif parte.startswith('**') and parte.endswith('**'):
            run = p.add_run(parte[2:-2]); run.bold = True
        elif parte.startswith('*') and parte.endswith('*'):
            run = p.add_run(parte[1:-1]); run.italic = True
        elif parte.startswith('`') and parte.endswith('`'):
            run = p.add_run(parte[1:-1]); run.font.name = 'Courier New'
            run.font.size = Pt(tamano - 1 if tamano else 10)
        else:
            run = p.add_run(parte)
        if tamano and run.font.size is None:
            run.font.size = Pt(tamano)

def parrafo(texto, tamano=None, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
            sangria_izq=None, sangria_francesa=False, antes=None, despues=None):
    p = doc.add_paragraph()
    p.alignment = alineacion
    if sangria_izq is not None:
        p.paragraph_format.left_indent = Cm(sangria_izq)
    if sangria_francesa:
        p.paragraph_format.first_line_indent = Cm(-0.75)
    if antes is not None:
        p.paragraph_format.space_before = Pt(antes)
    if despues is not None:
        p.paragraph_format.space_after = Pt(despues)
    agregar_runs(p, texto, tamano)
    return p

def encabezado(texto, nivel):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if nivel == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(limpiar(re.sub(r'[*`]', '', texto)))
    run.bold = True
    run.font.size = Pt(13 if nivel == 1 else 12)
    if nivel == 3:
        run.italic = True

def fila_tabla(linea):
    return [c.strip() for c in linea.strip().strip('|').split('|')]

def agregar_tabla(filas):
    ncols = len(filas[0])
    tabla = doc.add_table(rows=len(filas), cols=ncols)
    tabla.style = 'Table Grid'
    tabla.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, fila in enumerate(filas):
        for j, celda in enumerate(fila):
            if j >= ncols:
                continue
            c = tabla.cell(i, j)
            c.paragraphs[0].text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            agregar_runs(p, celda, tamano=10)
            if i == 0:
                for r in p.runs:
                    r.bold = True

# ---------------- Parseo ----------------
lineas = open(MD, encoding='utf-8').read().split('\n')
n = len(lineas)
i = 0
en_refs = False
primer_h1 = True
primer_h2 = True

while i < n:
    linea = lineas[i].rstrip()

    if not linea.strip() or linea.strip() == '---':
        i += 1
        continue

    # Titulo principal (ES) y titulo en ingles
    if linea.startswith('# ') and primer_h1:
        parrafo(linea[2:].strip(), tamano=14, alineacion=WD_ALIGN_PARAGRAPH.CENTER,
                antes=6, despues=4)
        for r in doc.paragraphs[-1].runs:
            r.bold = True
        primer_h1 = False
        i += 1
        continue
    if linea.startswith('## ') and primer_h2:
        p = parrafo(linea[3:].strip(), tamano=12, alineacion=WD_ALIGN_PARAGRAPH.CENTER,
                    despues=10)
        for r in p.runs:
            r.italic = True
        primer_h2 = False
        i += 1
        continue

    if linea.startswith('## '):
        h = linea[3:].strip()
        en_refs = h.startswith('References')
        encabezado(h, 1)
        i += 1
        continue
    if linea.startswith('### '):
        encabezado(linea[4:].strip(), 2)
        i += 1
        continue

    # Cita en bloque (pregunta de investigacion)
    if linea.startswith('> '):
        buf = []
        while i < n and lineas[i].startswith('>'):
            buf.append(lineas[i].lstrip('> ').strip())
            i += 1
        p = parrafo(' '.join(buf), sangria_izq=1.0, antes=6, despues=8)
        for r in p.runs:
            r.italic = True
        continue

    # Bloques de tabla / figura
    m = re.match(r'^\*\*(Table|Figure) (\d+)\*\*$', linea.strip())
    if m:
        tipo = m.group(1)
        parrafo(linea.strip(), antes=10, despues=1, alineacion=WD_ALIGN_PARAGRAPH.LEFT)
        i += 1
        parrafo(lineas[i].strip(), despues=6, alineacion=WD_ALIGN_PARAGRAPH.LEFT)
        i += 1
        while i < n and not lineas[i].strip():
            i += 1
        if tipo == 'Table':
            filas = []
            while i < n and lineas[i].strip().startswith('|'):
                if not re.match(r'^\|[\s:\-|]+\|$', lineas[i].strip()):
                    filas.append(fila_tabla(lineas[i]))
                i += 1
            agregar_tabla(filas)
        else:
            img = re.match(r'!\[[^\]]*\]\(([^)]+)\)', lineas[i].strip())
            ruta = os.path.normpath(os.path.join(DIR_PAPER, img.group(1)))
            doc.add_picture(ruta, width=Cm(14.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        while i < n and not lineas[i].strip():
            i += 1
        # Nota / Fuente
        if i < n and lineas[i].strip().startswith(('Nota:', 'Fuente:', '*Note.')):
            parrafo(lineas[i].strip(), tamano=10, antes=4, despues=10,
                    alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
        continue

    # Ecuacion display: $$latex$$ (N) -> ecuacion nativa OMML centrada con numero
    m_eq = re.match(r'^\$\$(.+)\$\$\s*\((\d+)\)\s*$', linea.strip())
    if m_eq:
        latex, num = m_eq.group(1).strip(), m_eq.group(2)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        p.add_run('\t')
        p._p.append(latex_a_omml(latex))
        p.add_run(f'\t({num})')
        i += 1
        continue

    # Bloque de autor (marcador @@AUTOR@@): centrado, bajo los títulos
    if linea.startswith('@@AUTOR@@'):
        p = parrafo(linea.replace('@@AUTOR@@', '').strip(),
                    alineacion=WD_ALIGN_PARAGRAPH.CENTER, despues=2)
        i += 1
        continue

    # Vinetas
    if linea.startswith('- '):
        txt = linea[2:].strip()
        i += 1
        while i < n and lineas[i].strip() and not lineas[i].startswith(('- ', '#', '|', '>')) \
                and not re.match(r'^\d+\. ', lineas[i]) \
                and not re.match(r'^\*\*(Table|Figure) \d+\*\*$', lineas[i].strip()):
            txt += ' ' + lineas[i].strip()
            i += 1
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        agregar_runs(p, txt)
        continue

    # Items numerados
    if re.match(r'^\d+\. ', linea):
        txt = linea
        i += 1
        while i < n and lineas[i].strip() and not lineas[i].startswith(('- ', '#', '|', '>')) \
                and not re.match(r'^\d+\. ', lineas[i]):
            txt += ' ' + lineas[i].strip()
            i += 1
        parrafo(txt, sangria_izq=0.6)
        continue

    # Parrafo normal (o referencia bibliografica)
    buf = [linea.strip()]
    i += 1
    while i < n and lineas[i].strip() and not lineas[i].startswith(('#', '- ', '|', '>')) \
            and not re.match(r'^\d+\. ', lineas[i]) \
            and not re.match(r'^\*\*(Table|Figure) \d+\*\*$', lineas[i].strip()):
        buf.append(lineas[i].strip())
        i += 1
    txt = ' '.join(buf)
    if en_refs:
        parrafo(txt, tamano=11, sangria_izq=0.75, sangria_francesa=True,
                despues=6, alineacion=WD_ALIGN_PARAGRAPH.LEFT)
    else:
        parrafo(txt)

doc.save(OUT)

# python-docx genera <w:zoom/> sin el atributo w:percent (obligatorio en el
# esquema OOXML). Se corrige para dejar el archivo formalmente valido.
import zipfile
import shutil
_tmp = OUT + '.tmp'
with zipfile.ZipFile(OUT, 'r') as zin, zipfile.ZipFile(_tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
    for item in zin.infolist():
        data = zin.read(item.filename)
        if item.filename == 'word/settings.xml':
            xml = data.decode('utf-8')
            xml = re.sub(r'<w:zoom(?![^>]*w:percent)([^>]*)/>',
                         r'<w:zoom w:percent="100"\1/>', xml)
            data = xml.encode('utf-8')
        zout.writestr(item, data)
shutil.move(_tmp, OUT)

print(f'OK: {OUT}')
print(f'Tamaño: {os.path.getsize(OUT):,} bytes')
