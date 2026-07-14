# -*- coding: utf-8 -*-
"""Genera los .docx de envio para Applied Economics (Taylor & Francis):
 - manuscrito principal (Times New Roman 12, doble espacio, margenes 2.5 cm,
   ecuaciones nativas OMML, tablas reemplazadas por marcadores [Table N near here])
 - archivo separado de tablas editables

Requiere: python-docx, latex2mathml, lxml; Office instalado (MML2OMML.XSL).
Uso:  python generar_docx_ae.py
"""
import os
import re
import sys
import zipfile
import shutil
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import latex2mathml.converter
from lxml import etree

AQUI = os.path.dirname(os.path.abspath(__file__))
_XSL = etree.XSLT(etree.parse(r'C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL'))

def latex_a_omml(latex):
    return _XSL(etree.fromstring(latex2mathml.converter.convert(latex))).getroot()

TOKEN = re.compile(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+?`)')

def limpiar(s):
    s = re.sub(r'!\[[^\]]*\]\([^)]*\)', '', s)
    s = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1', s)
    return s

def strip_md(s):
    return re.sub(r'[*`]', '', re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', s))

def construir(md_path, out_path, es_tablas=False):
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = (
        WD_LINE_SPACING.SINGLE if es_tablas else WD_LINE_SPACING.DOUBLE)
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Cm(2.5)

    def runs(p, texto, tam=None):
        for parte in TOKEN.split(limpiar(texto)):
            if not parte:
                continue
            if parte.startswith('***') and parte.endswith('***'):
                r = p.add_run(parte[3:-3]); r.bold = True; r.italic = True
            elif parte.startswith('**') and parte.endswith('**'):
                r = p.add_run(parte[2:-2]); r.bold = True
            elif parte.startswith('*') and parte.endswith('*'):
                r = p.add_run(parte[1:-1]); r.italic = True
            elif parte.startswith('`') and parte.endswith('`'):
                r = p.add_run(parte[1:-1]); r.font.name = 'Courier New'; r.font.size = Pt(10)
            else:
                r = p.add_run(parte)
            if tam and r.font.size is None:
                r.font.size = Pt(tam)

    def parrafo(texto, tam=None, al=WD_ALIGN_PARAGRAPH.JUSTIFY, izq=None, francesa=False,
                antes=None, despues=None):
        p = doc.add_paragraph(); p.alignment = al
        if izq is not None: p.paragraph_format.left_indent = Cm(izq)
        if francesa: p.paragraph_format.first_line_indent = Cm(-0.75)
        if antes is not None: p.paragraph_format.space_before = Pt(antes)
        if despues is not None: p.paragraph_format.space_after = Pt(despues)
        runs(p, texto, tam); return p

    def encabezado(texto, nivel):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14 if nivel == 1 else 10)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(strip_md(texto)); r.bold = True; r.font.size = Pt(13 if nivel == 1 else 12)
        if nivel == 3: r.italic = True

    def fila(l): return [c.strip() for c in l.strip().strip('|').split('|')]

    def tabla(filas):
        nc = len(filas[0])
        tb = doc.add_table(rows=len(filas), cols=nc); tb.style = 'Table Grid'
        tb.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, f in enumerate(filas):
            for j, celda in enumerate(f):
                if j >= nc: continue
                c = tb.cell(i, j); c.paragraphs[0].text = ''
                pp = c.paragraphs[0]; pp.paragraph_format.space_after = Pt(2)
                pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                runs(pp, celda, tam=10)
                if i == 0:
                    for rr in pp.runs: rr.bold = True

    lineas = open(md_path, encoding='utf-8').read().split('\n')
    n = len(lineas); i = 0
    en_refs = False; primer_h1 = True

    while i < n:
        linea = lineas[i].rstrip()
        if not linea.strip() or linea.strip() == '---':
            i += 1; continue

        if linea.startswith('@@AUTOR@@'):
            parrafo(linea.replace('@@AUTOR@@', '').strip(), al=WD_ALIGN_PARAGRAPH.CENTER, despues=2)
            i += 1; continue

        # ecuacion display $$latex$$ (N)
        meq = re.match(r'^\$\$(.+)\$\$\s*\((\d+)\)\s*$', linea.strip())
        if meq:
            p = doc.add_paragraph(); pf = p.paragraph_format
            pf.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
            pf.tab_stops.add_tab_stop(Cm(15.0), WD_TAB_ALIGNMENT.RIGHT)
            pf.space_before = Pt(6); pf.space_after = Pt(6)
            p.add_run('\t'); p._p.append(latex_a_omml(meq.group(1).strip()))
            p.add_run(f'\t({meq.group(2)})')
            i += 1; continue

        if linea.startswith('# ') and primer_h1:
            p = parrafo(linea[2:].strip(), tam=14, al=WD_ALIGN_PARAGRAPH.CENTER, antes=6, despues=6)
            for r in p.runs: r.bold = True
            primer_h1 = False; i += 1; continue

        if linea.startswith('## '):
            h = linea[3:].strip(); en_refs = h.startswith('References')
            encabezado(h, 1); i += 1; continue
        if linea.startswith('### '):
            encabezado(linea[4:].strip(), 2); i += 1; continue

        # marcador de tabla
        if re.match(r'^\[Table \d+ near here\]$', linea.strip()):
            p = parrafo(linea.strip(), al=WD_ALIGN_PARAGRAPH.CENTER, antes=6, despues=6)
            for r in p.runs: r.italic = True
            i += 1; continue

        # caption/tabla (solo en archivo de tablas)
        m = re.match(r'^\*\*Table (\d+)\*\*$', linea.strip())
        if m:
            parrafo(linea.strip(), antes=12, despues=1, al=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            parrafo(lineas[i].strip(), despues=6, al=WD_ALIGN_PARAGRAPH.LEFT); i += 1
            while i < n and not lineas[i].strip(): i += 1
            filas = []
            while i < n and lineas[i].strip().startswith('|'):
                if not re.match(r'^\|[\s:\-|]+\|$', lineas[i].strip()):
                    filas.append(fila(lineas[i]))
                i += 1
            tabla(filas)
            while i < n and not lineas[i].strip(): i += 1
            if i < n and lineas[i].strip().startswith('*Note.'):
                parrafo(lineas[i].strip(), tam=10, antes=4, despues=12, al=WD_ALIGN_PARAGRAPH.LEFT); i += 1
            continue

        # parrafo normal / referencia
        buf = [linea.strip()]; i += 1
        while i < n and lineas[i].strip() and not lineas[i].startswith(('#', '- ', '|', '>', '@@AUTOR@@')) \
                and not re.match(r'^\d+\. ', lineas[i]) \
                and not re.match(r'^\*\*Table \d+\*\*$', lineas[i].strip()) \
                and not re.match(r'^\$\$.+\$\$', lineas[i].strip()) \
                and not re.match(r'^\[Table \d+ near here\]$', lineas[i].strip()):
            buf.append(lineas[i].strip()); i += 1
        txt = ' '.join(buf)
        if en_refs:
            if txt.startswith('- '):
                txt = txt[2:]
            parrafo(txt, izq=0.75, francesa=True, despues=6, al=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            parrafo(txt)

    cp = doc.core_properties
    cp.author = 'Nicolás Guerrero Herrera'; cp.last_modified_by = ''
    cp.title = ('Same occupation, different pay: decomposing the gender wage gap '
                'within 4-digit occupations in Chile')
    doc.save(out_path)
    # correccion del atributo zoom (esquema OOXML)
    tmp = out_path + '.tmp'
    with zipfile.ZipFile(out_path) as zin, zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename == 'word/settings.xml':
                data = re.sub(r'<w:zoom(?![^>]*w:percent)([^>]*)/>',
                              r'<w:zoom w:percent="100"\1/>', data.decode('utf-8')).encode('utf-8')
            zout.writestr(it, data)
    shutil.move(tmp, out_path)
    print(f'OK: {os.path.basename(out_path)}  ({os.path.getsize(out_path):,} bytes)')

construir(os.path.join(AQUI, 'manuscrito_applied_economics.md'),
          os.path.join(AQUI, 'manuscrito_applied_economics.docx'), es_tablas=False)
construir(os.path.join(AQUI, 'tablas_applied_economics.md'),
          os.path.join(AQUI, 'tablas_applied_economics.docx'), es_tablas=True)
